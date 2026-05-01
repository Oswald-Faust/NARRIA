"""
narria.io.file_extractor — Extraction de texte depuis différents formats de fichiers.

Supporte :
  - .txt  : texte brut
  - .docx : Microsoft Word (via python-docx)
  - .pdf  : Adobe PDF (via pypdf)
  - .odt  : OpenDocument Text (via odfpy, optionnel)

Retourne le texte extrait et des métadonnées sur l'extraction (nombre de pages,
nombre de mots, avertissements éventuels sur la qualité de l'extraction).
"""

from __future__ import annotations

import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import Optional, List, Tuple


@dataclass
class ExtractionResult:
    """Résultat d'une extraction de texte depuis un fichier."""
    text: str                        # Texte extrait, nettoyé
    source_format: str               # 'txt' | 'docx' | 'pdf' | 'odt'
    source_filename: str
    word_count: int = 0
    char_count: int = 0
    page_count: Optional[int] = None     # pour PDF
    paragraph_count: Optional[int] = None # pour docx/odt
    warnings: List[str] = field(default_factory=list)
    title: str = ""                  # Titre extrait des métadonnées si disponible
    author: str = ""                 # Auteur extrait des métadonnées si disponible
    
    def to_dict(self) -> dict:
        return {
            'text': self.text,
            'source_format': self.source_format,
            'source_filename': self.source_filename,
            'word_count': self.word_count,
            'char_count': self.char_count,
            'page_count': self.page_count,
            'paragraph_count': self.paragraph_count,
            'warnings': self.warnings,
            'title': self.title,
            'author': self.author,
        }


class FileExtractor:
    """
    Extracteur de texte universel pour NARR'IA.
    
    Usage :
        extractor = FileExtractor()
        result = extractor.extract('/path/to/file.docx')
        if result.text:
            # Analyser result.text
            ...
    """
    
    SUPPORTED_FORMATS = {'.txt', '.docx', '.pdf', '.odt', '.epub'}
    
    def __init__(self, clean_text: bool = True):
        """
        Args:
            clean_text: si True, nettoie les artefacts typographiques (numéros
                        de page isolés, en-têtes/pieds récurrents, caractères
                        de contrôle, espaces multiples). Jamais ne touche au
                        contenu narratif lui-même.
        """
        self.clean_text = clean_text
    
    def extract(self, file_path: str | Path) -> ExtractionResult:
        """Extrait le texte d'un fichier selon son extension."""
        path = Path(file_path)
        if not path.exists():
            raise FileNotFoundError(f"Fichier introuvable : {path}")
        
        suffix = path.suffix.lower()
        if suffix not in self.SUPPORTED_FORMATS:
            raise ValueError(
                f"Format non supporté : {suffix}. "
                f"Formats acceptés : {', '.join(sorted(self.SUPPORTED_FORMATS))}"
            )
        
        if suffix == '.txt':
            result = self._extract_txt(path)
        elif suffix == '.docx':
            result = self._extract_docx(path)
        elif suffix == '.pdf':
            result = self._extract_pdf(path)
        elif suffix == '.odt':
            result = self._extract_odt(path)
        elif suffix == '.epub':
            result = self._extract_epub(path)
        else:
            raise ValueError(f"Format non géré : {suffix}")
        
        # Clean text if requested
        if self.clean_text:
            result.text = self._clean_text(result.text, result.source_format)
        
        # Update counts after cleaning
        result.char_count = len(result.text)
        result.word_count = len(result.text.split())
        
        return result
    
    # ───────────────────────────────────────────────────
    #  EXTRACTEURS SPÉCIFIQUES
    # ───────────────────────────────────────────────────
    
    def _extract_txt(self, path: Path) -> ExtractionResult:
        """Extrait le texte d'un fichier .txt (UTF-8, avec fallback Latin-1)."""
        try:
            text = path.read_text(encoding='utf-8')
        except UnicodeDecodeError:
            try:
                text = path.read_text(encoding='latin-1')
            except Exception as e:
                raise ValueError(f"Impossible de décoder le fichier texte : {e}")
        
        return ExtractionResult(
            text=text,
            source_format='txt',
            source_filename=path.name,
        )
    
    def _extract_docx(self, path: Path) -> ExtractionResult:
        """Extrait le texte d'un fichier .docx Word."""
        try:
            from docx import Document
        except ImportError:
            raise ImportError(
                "python-docx n'est pas installé. "
                "Installez-le avec : pip install python-docx"
            )
        
        try:
            doc = Document(str(path))
        except Exception as e:
            raise ValueError(f"Impossible d'ouvrir le fichier .docx : {e}")
        
        # Extract paragraphs
        paragraphs = []
        for p in doc.paragraphs:
            txt = p.text.strip()
            if txt:
                paragraphs.append(txt)
        
        # Also extract text from tables (preserves narrative content placed in tables)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        txt = p.text.strip()
                        if txt and txt not in paragraphs[-5:]:  # avoid short-window duplicates
                            paragraphs.append(txt)
        
        text = '\n\n'.join(paragraphs)
        
        # Extract metadata
        title = ""
        author = ""
        try:
            core_props = doc.core_properties
            title = core_props.title or ""
            author = core_props.author or ""
        except Exception:
            pass
        
        return ExtractionResult(
            text=text,
            source_format='docx',
            source_filename=path.name,
            paragraph_count=len(paragraphs),
            title=title,
            author=author,
        )
    
    def _extract_pdf(self, path: Path) -> ExtractionResult:
        """Extrait le texte d'un fichier .pdf."""
        try:
            import pypdf
        except ImportError:
            raise ImportError(
                "pypdf n'est pas installé. "
                "Installez-le avec : pip install pypdf"
            )
        
        warnings = []
        
        try:
            reader = pypdf.PdfReader(str(path))
        except Exception as e:
            raise ValueError(f"Impossible d'ouvrir le fichier .pdf : {e}")
        
        # Check if encrypted
        if reader.is_encrypted:
            try:
                reader.decrypt('')  # Try empty password
            except Exception:
                raise ValueError(
                    "Le fichier PDF est chiffré (protégé par mot de passe). "
                    "Veuillez le déverrouiller avant de le téléverser."
                )
        
        # Extract metadata
        title = ""
        author = ""
        try:
            meta = reader.metadata
            if meta:
                title = (meta.title or "")
                author = (meta.author or "")
        except Exception:
            pass
        
        # Extract text from each page
        page_texts = []
        for i, page in enumerate(reader.pages):
            try:
                page_text = page.extract_text() or ""
                page_texts.append(page_text)
            except Exception as e:
                warnings.append(f"Page {i+1} : erreur d'extraction ({str(e)[:80]})")
                page_texts.append("")
        
        text = '\n\n'.join(page_texts)
        
        # Warn if PDF appears to be scanned (very little text for many pages)
        n_pages = len(reader.pages)
        if n_pages >= 3 and len(text.strip()) < n_pages * 50:
            warnings.append(
                f"Ce PDF semble être un scan (très peu de texte détecté pour "
                f"{n_pages} pages). Pour les PDFs scannés, une OCR est nécessaire, "
                f"ce que NARR'IA ne propose pas actuellement. Si votre document "
                f"est un vrai PDF avec du texte, vérifiez qu'il n'est pas "
                f"excessivement protégé."
            )
        
        return ExtractionResult(
            text=text,
            source_format='pdf',
            source_filename=path.name,
            page_count=n_pages,
            title=title,
            author=author,
            warnings=warnings,
        )
    
    def _extract_odt(self, path: Path) -> ExtractionResult:
        """Extrait le texte d'un fichier .odt OpenDocument."""
        try:
            from odf.opendocument import load
            from odf import text, teletype
        except ImportError:
            raise ImportError(
                "odfpy n'est pas installé. "
                "Installez-le avec : pip install odfpy"
            )
        
        try:
            doc = load(str(path))
        except Exception as e:
            raise ValueError(f"Impossible d'ouvrir le fichier .odt : {e}")
        
        # Extract all paragraphs
        paragraphs = []
        for p in doc.getElementsByType(text.P):
            t = teletype.extractText(p).strip()
            if t:
                paragraphs.append(t)
        
        text_str = '\n\n'.join(paragraphs)
        
        # Try metadata
        title = ""
        author = ""
        try:
            meta = doc.meta
            if meta:
                title_el = meta.getElementsByType(type('Title', (), {'qname': ('urn:oasis:names:tc:opendocument:xmlns:office:1.0', 'title')}))
                # This is tricky with odfpy; we'll skip for simplicity
        except Exception:
            pass
        
        return ExtractionResult(
            text=text_str,
            source_format='odt',
            source_filename=path.name,
            paragraph_count=len(paragraphs),
            title=title,
            author=author,
        )
    
    def _extract_epub(self, path: Path) -> ExtractionResult:
        """
        Extrait le texte d'un fichier .epub (livre numérique).
        
        Parcourt les chapitres dans l'ordre de la spine (ordre de lecture),
        extrait le HTML de chaque chapitre et en convertit le contenu en
        texte brut via BeautifulSoup.
        """
        try:
            import ebooklib
            from ebooklib import epub
            from bs4 import BeautifulSoup
        except ImportError as e:
            raise ImportError(
                "Les bibliothèques ebooklib et beautifulsoup4 ne sont pas "
                "installées. Installez-les avec : "
                "pip install ebooklib beautifulsoup4"
            )
        
        warnings = []
        
        try:
            book = epub.read_epub(str(path))
        except Exception as e:
            raise ValueError(f"Impossible d'ouvrir le fichier .epub : {e}")
        
        # Extract metadata
        title = ""
        author = ""
        try:
            title_meta = book.get_metadata('DC', 'title')
            if title_meta:
                title = title_meta[0][0] or ""
            creator_meta = book.get_metadata('DC', 'creator')
            if creator_meta:
                author = creator_meta[0][0] or ""
        except Exception:
            pass
        
        # Walk the spine (reading order) to collect chapters
        chapter_texts = []
        n_chapters = 0
        
        for item_id, _linear in book.spine:
            item = book.get_item_with_id(item_id)
            if item is None:
                continue
            if item.get_type() != ebooklib.ITEM_DOCUMENT:
                continue
            
            try:
                html_content = item.get_content()
                soup = BeautifulSoup(html_content, 'html.parser')
                
                # Remove <script> and <style> blocks
                for tag in soup(['script', 'style', 'nav', 'header', 'footer']):
                    tag.decompose()
                
                # Extract text with paragraph separation
                # Replace <br> with newlines
                for br in soup.find_all('br'):
                    br.replace_with('\n')
                
                # For block-level elements, ensure they're on separate lines
                block_tags = ['p', 'div', 'h1', 'h2', 'h3', 'h4', 'h5', 'h6',
                              'li', 'blockquote', 'pre']
                chapter_parts = []
                for el in soup.find_all(block_tags):
                    txt = el.get_text(strip=True)
                    if txt:
                        chapter_parts.append(txt)
                
                # Fallback: if no block-level content found, get all text
                if not chapter_parts:
                    text_content = soup.get_text(separator='\n', strip=True)
                    if text_content:
                        chapter_parts.append(text_content)
                
                chapter_text = '\n\n'.join(chapter_parts)
                if chapter_text.strip():
                    chapter_texts.append(chapter_text)
                    n_chapters += 1
            except Exception as e:
                warnings.append(f"Chapitre {n_chapters+1} : erreur d'extraction ({str(e)[:80]})")
        
        text = '\n\n'.join(chapter_texts)
        
        if not text.strip():
            warnings.append(
                "Aucun texte n'a pu être extrait de ce fichier EPUB. "
                "Il est possible que le livre soit protégé par DRM ou que "
                "sa structure soit non standard."
            )
        
        return ExtractionResult(
            text=text,
            source_format='epub',
            source_filename=path.name,
            paragraph_count=n_chapters,  # We track chapters rather than paragraphs for EPUB
            title=title,
            author=author,
            warnings=warnings,
        )
    
    # ───────────────────────────────────────────────────
    #  NETTOYAGE DU TEXTE
    # ───────────────────────────────────────────────────
    
    def _clean_text(self, text: str, source_format: str) -> str:
        """
        Nettoie le texte extrait :
        - retire les lignes isolées contenant uniquement un numéro (numéros de page)
        - retire les headers/footers récurrents (détectés par répétition)
        - normalise les espaces multiples et les retours à la ligne excessifs
        - retire les caractères de contrôle
        """
        if not text:
            return text
        
        # Normalize line endings
        text = re.sub(r'\r\n?', '\n', text)
        
        # Remove control chars except \n and \t
        text = re.sub(r'[\x00-\x08\x0B-\x1F\x7F]', '', text)
        
        # For PDFs specifically: detect and remove repeated headers/footers
        if source_format == 'pdf':
            text = self._remove_recurring_lines(text)
            text = self._remove_page_numbers(text)
        
        # Collapse multiple spaces (but preserve line breaks)
        text = re.sub(r'[ \t]+', ' ', text)
        
        # Collapse multiple blank lines to at most 2
        text = re.sub(r'\n{3,}', '\n\n', text)
        
        # Strip leading/trailing whitespace
        text = text.strip()
        
        return text
    
    def _remove_recurring_lines(self, text: str) -> str:
        """
        Détecte les lignes qui se répètent suspicieusement (probablement des
        headers/footers de page) et les retire.
        
        Heuristique : une ligne qui apparaît plus de 3 fois ET qui fait moins
        de 80 caractères est probablement un header/footer.
        """
        from collections import Counter
        lines = text.split('\n')
        
        # Count non-empty lines
        line_counts = Counter(line.strip() for line in lines if line.strip())
        
        # Lines that look like recurring headers/footers
        recurring = {
            line for line, count in line_counts.items()
            if count >= 3 and 0 < len(line) < 80 and not line[0].islower()
        }
        
        if not recurring:
            return text
        
        # Filter them out
        filtered = []
        for line in lines:
            if line.strip() not in recurring:
                filtered.append(line)
        
        return '\n'.join(filtered)
    
    def _remove_page_numbers(self, text: str) -> str:
        """Retire les lignes isolées contenant uniquement un nombre (numéros de page)."""
        lines = text.split('\n')
        filtered = []
        for line in lines:
            stripped = line.strip()
            # Skip lines that are only page numbers (1-4 digits)
            if re.fullmatch(r'\d{1,4}', stripped):
                continue
            # Skip lines like "- 12 -" or "12."
            if re.fullmatch(r'[\-\s]*\d{1,4}[\-\s.]*', stripped):
                continue
            filtered.append(line)
        return '\n'.join(filtered)


# Convenience function
def extract_text(file_path: str | Path, clean: bool = True) -> ExtractionResult:
    """Fonction utilitaire : extrait le texte d'un fichier en une ligne."""
    return FileExtractor(clean_text=clean).extract(file_path)
